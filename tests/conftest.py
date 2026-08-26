"""Shared fixtures. Binary formats (pdf/docx/epub) are generated on the fly with
PyMuPDF/python-docx/ebooklib rather than committed as binary test fixtures."""

from __future__ import annotations

from pathlib import Path

import pytest


@pytest.fixture
def md_file(tmp_path: Path) -> Path:
    content = """# Attention Is All You Need

## 1 Introduction

The dominant sequence transduction models are based on complex recurrent networks.

## 2 Background

Some background text about **self-attention** and prior work.

### 2.1 Self Attention

Self-attention relates different positions of a single sequence.

- First point
- Second point
  - Nested point

```python
class MultiHeadAttention:
    pass
```

> Attention is all you need.

| Model | BLEU |
| --- | --- |
| Transformer | 28.4 |
| ByteNet | 23.75 |
"""
    path = tmp_path / "fixture.md"
    path.write_text(content)
    return path


@pytest.fixture
def txt_file(tmp_path: Path) -> Path:
    content = (
        "Chapter One\n\n"
        "This is the first paragraph of a plain text book.\n\n"
        "This is the second paragraph, discussing attention mechanisms in detail.\n"
    )
    path = tmp_path / "fixture.txt"
    path.write_text(content)
    return path


@pytest.fixture
def docx_file(tmp_path: Path) -> Path:
    import docx

    d = docx.Document()
    d.add_heading("Designing Data-Intensive Applications", level=0)
    d.add_heading("Chapter 5: Replication", level=1)
    d.add_paragraph("Replication means keeping a copy of the same data on multiple machines.")
    d.add_heading("Leaderless Replication", level=2)
    d.add_paragraph("Some databases abandon the concept of a leader entirely.")
    d.add_paragraph("First bullet", style="List Bullet")
    d.add_paragraph("Second bullet", style="List Bullet")
    table = d.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Model"
    table.rows[0].cells[1].text = "BLEU"
    table.rows[1].cells[0].text = "Transformer"
    table.rows[1].cells[1].text = "28.4"

    path = tmp_path / "fixture.docx"
    d.save(path)
    return path


@pytest.fixture
def pdf_file(tmp_path: Path) -> Path:
    import pymupdf

    pdf = pymupdf.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "Attention Is All You Need", fontsize=24)
    page.insert_text((72, 110), "3.2 Multi-Head Attention", fontsize=16)
    page.insert_text((72, 140), "Instead of performing a single attention function with", fontsize=10)
    page.insert_text((72, 155), "d_model-dimensional keys, values and queries, we found it", fontsize=10)
    page.insert_text((72, 170), "beneficial to linearly project the queries h times.", fontsize=10)
    page2 = pdf.new_page()
    page2.insert_text((72, 72), "3.3 Masked Attention", fontsize=16)
    page2.insert_text((72, 100), "The decoder applies masking to prevent positions from attending", fontsize=10)
    page2.insert_text((72, 115), "to subsequent positions.", fontsize=10)

    path = tmp_path / "fixture.pdf"
    pdf.save(path)
    return path


@pytest.fixture
def pdf_with_duplicate_outline_pages(tmp_path: Path) -> Path:
    """A PDF outline (bookmarks) where multiple entries share the same page — common for
    front-matter sections (e.g. "Prerequisites" and "Acknowledgments" both on page 1)."""
    import pymupdf

    pdf = pymupdf.open()
    p1 = pdf.new_page()
    p1.insert_text((72, 72), "Front Matter", fontsize=20)
    p2 = pdf.new_page()
    p2.insert_text((72, 72), "Chapter One", fontsize=20)
    p2.insert_text((72, 110), "Some content on chapter one page.", fontsize=10)

    pdf.set_toc(
        [
            [1, "Prerequisites", 1],
            [1, "Acknowledgments", 1],
            [1, "Chapter One", 2],
            [2, "Section 1.1", 2],
        ]
    )
    path = tmp_path / "dup_page.pdf"
    pdf.save(path)
    return path


@pytest.fixture
def epub_file(tmp_path: Path) -> Path:
    from ebooklib import epub

    book = epub.EpubBook()
    book.set_identifier("id123")
    book.set_title("Test Book")
    book.set_language("en")
    book.add_author("Jane Doe")

    c1 = epub.EpubHtml(title="Chapter 1", file_name="chap1.xhtml", lang="en")
    c1.content = "<h1>Chapter 1</h1><p>This is the first chapter about attention.</p><ul><li>point a</li><li>point b</li></ul>"
    c2 = epub.EpubHtml(title="Chapter 2", file_name="chap2.xhtml", lang="en")
    c2.content = "<p>This chapter has no heading tag at all, just body text.</p>"

    book.add_item(c1)
    book.add_item(c2)
    book.toc = (epub.Link("chap1.xhtml", "Chapter 1", "c1"), epub.Link("chap2.xhtml", "Chapter 2", "c2"))
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = ["nav", c1, c2]

    path = tmp_path / "fixture.epub"
    epub.write_epub(str(path), book)
    return path


@pytest.fixture
def big_md_file(tmp_path: Path) -> Path:
    lines = ["# Big Test Document", ""]
    for i in range(1, 61):
        lines.append(f"## Section {i}")
        lines.append("")
        for j in range(1, 6):
            lines.append(
                f"This is paragraph {j} of section {i}, discussing attention mechanisms "
                "and transformers in enough detail to wrap across a couple of lines."
            )
            lines.append("")
    path = tmp_path / "big.md"
    path.write_text("\n".join(lines))
    return path


@pytest.fixture
def db_path(tmp_path: Path) -> Path:
    return tmp_path / "pinax.db"


@pytest.fixture
def isolated_home(tmp_path: Path, monkeypatch) -> Path:
    """Point platformdirs-based config/data/cache dirs at a throwaway directory so app-level
    tests never touch (or get polluted by) the real user's pinax config/database."""
    home = tmp_path / "home"
    home.mkdir()
    monkeypatch.setenv("HOME", str(home))
    return home
